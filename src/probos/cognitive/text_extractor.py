"""AD-720d (Wave 139): non-image attachment text extraction.

AD-720a-1 (Wave 162) adds PDF / DOCX / XLSX extraction via permissive deps
(pypdf BSD-3, python-docx MIT, openpyxl MIT). The PDF/DOCX/XLSX path is
gated by ``AttachmentsConfig.pdf_extraction_enabled`` (default-OFF
transitional flag).
"""

from __future__ import annotations

import json
import logging

logger = logging.getLogger(__name__)


_TEXT_PASSTHROUGH_MIMES: frozenset[str] = frozenset({
    "text/plain",
    "text/markdown",
    "text/csv",
})


_MAX_PDF_PAGES = 100
_MAX_XLSX_ROWS = 10_000


def _extract_pdf(blob: bytes, max_bytes: int) -> tuple[str, bool]:
    """AD-720a-1: page-by-page PDF extraction, capped at 100 pages."""
    import io

    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(blob))
    pages = reader.pages[:_MAX_PDF_PAGES]
    page_truncated = len(reader.pages) > _MAX_PDF_PAGES

    parts: list[str] = []
    running_bytes = 0
    byte_truncated = False
    for idx, page in enumerate(pages):
        try:
            page_text = page.extract_text() or ""
        except Exception:
            logger.warning(
                "AD-720a-1: pypdf failed on page %d; emitting empty page text",
                idx, exc_info=True,
            )
            page_text = ""
        chunk = f"\n--- Page {idx + 1} ---\n{page_text}"
        encoded = chunk.encode("utf-8")
        if running_bytes + len(encoded) > max_bytes:
            remaining = max_bytes - running_bytes
            if remaining > 0:
                parts.append(encoded[:remaining].decode("utf-8", errors="ignore"))
            byte_truncated = True
            break
        parts.append(chunk)
        running_bytes += len(encoded)

    text = "".join(parts).lstrip("\n")
    truncated = page_truncated or byte_truncated
    if truncated:
        text = text + "\n[TRUNCATED]"
    return (text, truncated)


def _extract_docx(blob: bytes, max_bytes: int) -> tuple[str, bool]:
    """AD-720a-1: DOCX extraction — paragraphs + table cells, in document order."""
    import io

    from docx import Document

    doc = Document(io.BytesIO(blob))

    parts: list[str] = []
    running_bytes = 0
    truncated = False

    for para in doc.paragraphs:
        text = para.text
        if not text:
            continue
        encoded = (text + "\n").encode("utf-8")
        if running_bytes + len(encoded) > max_bytes:
            remaining = max_bytes - running_bytes
            if remaining > 0:
                parts.append(encoded[:remaining].decode("utf-8", errors="ignore"))
            truncated = True
            break
        parts.append(text + "\n")
        running_bytes += len(encoded)

    if not truncated:
        for table in doc.tables:
            row_truncated = False
            for row in table.rows:
                row_text = " | ".join(cell.text for cell in row.cells) + "\n"
                encoded = row_text.encode("utf-8")
                if running_bytes + len(encoded) > max_bytes:
                    remaining = max_bytes - running_bytes
                    if remaining > 0:
                        parts.append(
                            encoded[:remaining].decode("utf-8", errors="ignore")
                        )
                    truncated = True
                    row_truncated = True
                    break
                parts.append(row_text)
                running_bytes += len(encoded)
            if row_truncated:
                break

    text = "".join(parts)
    if truncated:
        text = text + "\n[TRUNCATED]"
    return (text, truncated)


def _extract_xlsx(blob: bytes, max_bytes: int) -> tuple[str, bool]:
    """AD-720a-1: XLSX extraction — sheet-by-sheet, cell-by-cell, max 10k rows total."""
    import io

    from openpyxl import load_workbook

    wb = load_workbook(io.BytesIO(blob), data_only=True, read_only=True)

    parts: list[str] = []
    running_bytes = 0
    rows_emitted = 0
    truncated = False

    try:
        for sheet_name in wb.sheetnames:
            if truncated:
                break
            ws = wb[sheet_name]
            header = f"\n=== Sheet: {sheet_name} ===\n"
            parts.append(header)
            running_bytes += len(header.encode("utf-8"))

            for row in ws.iter_rows(values_only=True):
                if rows_emitted >= _MAX_XLSX_ROWS:
                    truncated = True
                    break
                row_text = " | ".join(
                    "" if v is None else str(v) for v in row
                ) + "\n"
                encoded = row_text.encode("utf-8")
                if running_bytes + len(encoded) > max_bytes:
                    remaining = max_bytes - running_bytes
                    if remaining > 0:
                        parts.append(
                            encoded[:remaining].decode("utf-8", errors="ignore")
                        )
                    truncated = True
                    break
                parts.append(row_text)
                running_bytes += len(encoded)
                rows_emitted += 1
    finally:
        wb.close()

    text = "".join(parts).lstrip("\n")
    if truncated:
        text = text + "\n[TRUNCATED]"
    return (text, truncated)


_DOCUMENT_DISPATCH = {
    "application/pdf": _extract_pdf,
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": _extract_docx,
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": _extract_xlsx,
}


async def extract_text(
    blob: bytes,
    mime: str,
    *,
    max_bytes: int,
) -> tuple[str, bool]:
    """Return ``(extracted_text, was_truncated)`` for non-image attachment bytes.

    AD-720a-1 document MIMEs (PDF/DOCX/XLSX) dispatch to dedicated helpers;
    parser exceptions bubble as ValueError. AD-720d MIMEs (text/plain,
    text/markdown, text/csv, application/json) unchanged.
    """
    if mime in _DOCUMENT_DISPATCH:
        try:
            return _DOCUMENT_DISPATCH[mime](blob, max_bytes)
        except Exception as exc:
            raise ValueError(
                f"AD-720a-1: failed to extract text from {mime!r}: {exc}"
            ) from exc

    if mime in _TEXT_PASSTHROUGH_MIMES:
        text = blob.decode("utf-8", errors="strict")
    elif mime == "application/json":
        parsed = json.loads(blob.decode("utf-8", errors="strict"))
        text = json.dumps(parsed, indent=2)
    else:
        raise ValueError(f"unsupported MIME for text extraction: {mime!r}")

    encoded = text.encode("utf-8")
    if len(encoded) > max_bytes:
        truncated = encoded[:max_bytes].decode("utf-8", errors="ignore")
        return (truncated + "\n[TRUNCATED]", True)
    return (text, False)
