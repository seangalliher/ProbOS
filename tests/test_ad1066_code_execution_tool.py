"""AD-1066: CodeExecutionTool — sandboxed Python with artifact capture.

Substrate-honest integration: a REAL ``SubprocessSandbox`` runs actual Python and
a REAL ``ArtifactStore`` + a Protocol-faithful ``AttachmentStore`` capture the
produced files. Proves the Claude Cowork / Codex parity loop — run code →
produced file → downloadable artifact on the chat thread.
"""

from __future__ import annotations

from pathlib import Path
from types import SimpleNamespace

from probos.artifacts import ArtifactStore
from probos.config import ExecutionConfig
from probos.tools.code_execution_tool import CodeExecutionTool


class _FakeAttachmentStore:
    """Protocol-faithful in-memory AttachmentStore (AD-720 write signature:
    positional content_hash/blob/mime, keyword-only origin)."""

    def __init__(self) -> None:
        self.blobs: dict[str, tuple[bytes, str, str]] = {}

    async def write(self, content_hash, blob, mime, *, origin="chat_attachment"):
        self.blobs[content_hash] = (blob, mime, origin)
        return Path(f"/fake/{content_hash}")


def _runtime(tmp_path: Path, *, enabled: bool = True):
    return SimpleNamespace(
        config=SimpleNamespace(
            execution=ExecutionConfig(
                enabled=enabled, scratch_dir=str(tmp_path / "scratch"),
            ),
        ),
        artifact_store=ArtifactStore(tmp_path / "artifacts.db"),
        attachment_store=_FakeAttachmentStore(),
    )


def _ctx(thread_id="thread-1", agent_id="ezri"):
    return {"thread_id": thread_id, "agent_id": agent_id}


async def test_runs_code_and_captures_produced_file_as_artifact(tmp_path) -> None:
    runtime = _runtime(tmp_path)
    tool = CodeExecutionTool(runtime=runtime)

    code = (
        "with open('report.txt', 'w', encoding='utf-8') as f:\n"
        "    f.write('Recommendations for ProbOS\\n')\n"
        "print('document written')\n"
    )
    result = await tool.invoke({"code": code}, _ctx())

    assert result.error is None
    assert result.output["success"] is True
    assert "document written" in result.output["stdout"]
    # The produced file became a downloadable artifact on the thread.
    assert result.output["artifacts"] == ["report.txt"]
    arts = runtime.artifact_store.list_thread_latest("thread-1")
    assert len(arts) == 1
    assert arts[0].name == "report.txt"
    assert arts[0].created_by == "ezri"
    assert arts[0].mime == "text/plain"
    # The bytes landed in the attachment store under the artifact's hash.
    assert arts[0].content_hash in runtime.attachment_store.blobs
    blob, _mime, origin = runtime.attachment_store.blobs[arts[0].content_hash]
    assert b"Recommendations for ProbOS" in blob
    assert origin == "agent_artifact"


async def test_does_not_capture_the_sandbox_script(tmp_path) -> None:
    runtime = _runtime(tmp_path)
    tool = CodeExecutionTool(runtime=runtime)
    result = await tool.invoke({"code": "print('no files produced')"}, _ctx())
    assert result.error is None
    # script.py is the sandbox's own file — never surfaced as a deliverable.
    assert result.output["artifacts"] == []
    assert runtime.artifact_store.list_thread_latest("thread-1") == []


async def test_docx_extension_gets_word_mime(tmp_path) -> None:
    runtime = _runtime(tmp_path)
    tool = CodeExecutionTool(runtime=runtime)
    code = "open('out.docx','wb').write(b'PK\\x03\\x04stub-docx-bytes')"
    result = await tool.invoke({"code": code}, _ctx())
    arts = runtime.artifact_store.list_thread_latest("thread-1")
    assert len(arts) == 1
    assert arts[0].mime == (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


async def test_disabled_returns_error(tmp_path) -> None:
    runtime = _runtime(tmp_path, enabled=False)
    tool = CodeExecutionTool(runtime=runtime)
    result = await tool.invoke({"code": "print(1)"}, _ctx())
    assert result.error is not None
    assert "disabled" in result.error


async def test_no_code_returns_error(tmp_path) -> None:
    runtime = _runtime(tmp_path)
    tool = CodeExecutionTool(runtime=runtime)
    result = await tool.invoke({"code": "   "}, _ctx())
    assert result.error is not None


async def test_no_thread_id_runs_but_captures_nothing(tmp_path) -> None:
    runtime = _runtime(tmp_path)
    tool = CodeExecutionTool(runtime=runtime)
    code = "open('x.txt','w').write('hi')\nprint('ran')"
    result = await tool.invoke({"code": code}, {"agent_id": "ezri"})  # no thread_id
    assert result.error is None
    assert "ran" in result.output["stdout"]
    assert result.output["artifacts"] == []


def test_tool_protocol_shape() -> None:
    tool = CodeExecutionTool(runtime=SimpleNamespace())
    assert tool.tool_id == "run_python"
    assert isinstance(tool.input_schema, dict)
    assert "code" in tool.input_schema["properties"]
    assert "document" in tool.description.lower()


# ── AD-1066 wiring: the AgenticLoop offers run_python + captures its output ──
# End-to-end through WorkItemAgenticExecutor (the shared loop) with a scripted
# LLM, proving: registered-when-enabled + offered-to-the-model + invoked + the
# thread_id threaded all the way to the produced artifact.


class _FakeLLMResponse:
    def __init__(self, content_blocks, content="", tokens=1):
        self.content_blocks = content_blocks
        self.content = content
        self.tokens_used = tokens


class _ScriptedLLM:
    def __init__(self, responses):
        self._responses = list(responses)
        self.last_tools = None

    async def complete(self, req, **_kwargs):
        self.last_tools = list(getattr(req, "tools", None) or [])
        if self._responses:
            return self._responses.pop(0)
        return _FakeLLMResponse(content_blocks=[], content="done")


def _tool_use(tool_id, args):
    from probos.cognitive.swe_harness.tool_call import ToolCallRequest, ToolUseBlock

    return _FakeLLMResponse(
        content_blocks=[ToolUseBlock(tool_call=ToolCallRequest(name=tool_id, arguments=args))],
        content="",
        tokens=1,
    )


def _text(text):
    return _FakeLLMResponse(content_blocks=[], content=text, tokens=1)


def _tool_names(last_tools):
    return [
        (t.get("function", {}) or {}).get("name") or t.get("name")
        for t in (last_tools or [])
    ]


def _loop_runtime(tmp_path, *, enabled: bool):
    from probos.tools.permissions import ToolPermissionStore
    from probos.tools.registry import ToolRegistry

    return SimpleNamespace(
        config=SimpleNamespace(
            execution=ExecutionConfig(
                enabled=enabled, scratch_dir=str(tmp_path / "scratch"),
            ),
            mcp=None,
        ),
        tool_registry=ToolRegistry(),
        tool_permission_store=ToolPermissionStore(),
        intent_bus=None,
        intent_grant_store=None,
        mcp_workbench=None,
        attachment_store=_FakeAttachmentStore(),
        artifact_store=ArtifactStore(tmp_path / "artifacts.db"),
        emit_event=None,
    )


async def test_loop_offers_run_python_and_captures_output_as_artifact(tmp_path) -> None:
    from probos.cognitive.agentic_dispatch import WorkItemAgenticExecutor

    runtime = _loop_runtime(tmp_path, enabled=True)
    code = (
        "with open('summary.txt', 'w', encoding='utf-8') as f:\n"
        "    f.write('crew agent deliverable')\n"
        "print('wrote summary.txt')\n"
    )
    llm = _ScriptedLLM([
        _tool_use("run_python", {"code": code}),
        _text("I've prepared summary.txt for you, Captain."),
    ])
    executor = WorkItemAgenticExecutor(llm_client=llm)

    outcome = await executor.run(
        agent_id="ezri",
        instructions="You are Ezri.",
        task_text="make me a summary file",
        runtime=runtime,
        department="counseling",
        rank="lieutenant",
        thread_id="thread-xyz",
        max_iterations=4,
        tier="standard",
    )

    assert outcome.final_text == "I've prepared summary.txt for you, Captain."
    # run_python was registered + offered to the model.
    assert "run_python" in _tool_names(llm.last_tools)
    assert runtime.tool_registry.get("run_python") is not None
    # The produced file is a downloadable artifact on the threaded thread_id.
    arts = runtime.artifact_store.list_thread_latest("thread-xyz")
    assert len(arts) == 1
    assert arts[0].name == "summary.txt"
    assert arts[0].created_by == "ezri"
    assert arts[0].content_hash in runtime.attachment_store.blobs


async def test_loop_omits_run_python_when_execution_disabled(tmp_path) -> None:
    from probos.cognitive.agentic_dispatch import WorkItemAgenticExecutor

    runtime = _loop_runtime(tmp_path, enabled=False)
    llm = _ScriptedLLM([_text("hello, Captain")])
    executor = WorkItemAgenticExecutor(llm_client=llm)

    outcome = await executor.run(
        agent_id="ezri",
        instructions="You are Ezri.",
        task_text="hi",
        runtime=runtime,
        thread_id="t",
        max_iterations=2,
    )

    assert outcome.final_text == "hello, Captain"
    assert "run_python" not in _tool_names(llm.last_tools)
    assert runtime.tool_registry.get("run_python") is None


async def test_loop_produces_a_real_docx_via_python_docx(tmp_path) -> None:
    """THE GOAL (Cowork parity): an agent in the conversational loop produces a
    REAL, re-openable Word document (.docx) that lands as a downloadable artifact
    on the chat thread. python-docx is a core ProbOS dependency, so this runs the
    actual generation path end-to-end — only a live LLM + dm_agentic.enabled +
    execution.enabled separate this from the Captain's real chat experience."""
    import io

    from docx import Document

    from probos.cognitive.agentic_dispatch import WorkItemAgenticExecutor

    runtime = _loop_runtime(tmp_path, enabled=True)
    code = (
        "from docx import Document\n"
        "doc = Document()\n"
        "doc.add_heading('ProbOS Recommendations', 0)\n"
        "doc.add_paragraph('Prepared by Ezri, Ship\\'s Counselor.')\n"
        "doc.add_paragraph('1. Rolling cognitive baselines.')\n"
        "doc.save('recommendations.docx')\n"
        "print('docx written')\n"
    )
    llm = _ScriptedLLM([
        _tool_use("run_python", {"code": code}),
        _text("Done — recommendations.docx is ready to download, Captain."),
    ])
    executor = WorkItemAgenticExecutor(llm_client=llm)

    outcome = await executor.run(
        agent_id="ezri",
        instructions="You are Ezri, Ship's Counselor.",
        task_text="write up your recommendations as a Word document",
        runtime=runtime,
        thread_id="thread-docx",
        max_iterations=4,
    )

    assert outcome.final_text == "Done — recommendations.docx is ready to download, Captain."
    arts = runtime.artifact_store.list_thread_latest("thread-docx")
    assert len(arts) == 1
    assert arts[0].name == "recommendations.docx"
    assert arts[0].mime == (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    blob, _mime, _origin = runtime.attachment_store.blobs[arts[0].content_hash]
    # A real .docx is a ZIP archive (PK magic) ...
    assert blob[:2] == b"PK"
    # ... that python-docx can re-open with the authored content intact.
    reopened = Document(io.BytesIO(blob))
    texts = "\n".join(p.text for p in reopened.paragraphs)
    assert "ProbOS Recommendations" in texts
    assert "Rolling cognitive baselines" in texts
