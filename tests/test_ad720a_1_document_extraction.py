"""AD-720a-1: PDF / DOCX / XLSX document text extraction tests."""
from __future__ import annotations

import asyncio
import io

import pytest

from probos.cognitive.text_extractor import (
    _MAX_PDF_PAGES,
    _MAX_XLSX_ROWS,
    extract_text,
)
from probos.config import SystemConfig


# ── Fixture builders ────────────────────────────────────────────


def _make_pdf(num_pages: int) -> bytes:
    from pypdf import PdfWriter
    w = PdfWriter()
    for _ in range(num_pages):
        w.add_blank_page(width=200, height=200)
    buf = io.BytesIO()
    w.write(buf)
    return buf.getvalue()


def _make_docx(paragraphs: list[str], table_rows: list[list[str]] | None = None) -> bytes:
    from docx import Document
    doc = Document()
    for p in paragraphs:
        doc.add_paragraph(p)
    if table_rows:
        cols = len(table_rows[0])
        table = doc.add_table(rows=len(table_rows), cols=cols)
        for r_idx, row in enumerate(table_rows):
            for c_idx, val in enumerate(row):
                table.rows[r_idx].cells[c_idx].text = val
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _make_xlsx(sheets: dict[str, list[list[str]]]) -> bytes:
    from openpyxl import Workbook
    wb = Workbook()
    wb.remove(wb.active)
    for name, rows in sheets.items():
        ws = wb.create_sheet(name)
        for row in rows:
            ws.append(row)
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


# ── PDF tests ──────────────────────────────────────────────────


@pytest.mark.asyncio
async def test_pdf_happy_path() -> None:
    blob = _make_pdf(3)
    text, truncated = await extract_text(blob, "application/pdf", max_bytes=64_000)
    assert "--- Page 1 ---" in text
    assert "--- Page 2 ---" in text
    assert "--- Page 3 ---" in text
    assert truncated is False


@pytest.mark.asyncio
async def test_pdf_page_cap() -> None:
    blob = _make_pdf(_MAX_PDF_PAGES + 5)
    text, truncated = await extract_text(blob, "application/pdf", max_bytes=10_000_000)
    assert truncated is True
    assert "[TRUNCATED]" in text
    assert f"--- Page {_MAX_PDF_PAGES} ---" in text


@pytest.mark.asyncio
async def test_pdf_byte_cap() -> None:
    blob = _make_pdf(10)
    text, truncated = await extract_text(blob, "application/pdf", max_bytes=80)
    assert truncated is True
    assert "[TRUNCATED]" in text


@pytest.mark.asyncio
async def test_pdf_corrupt_bytes_raises() -> None:
    with pytest.raises(ValueError, match="AD-720a-1"):
        await extract_text(b"this is not a pdf", "application/pdf", max_bytes=64_000)


# ── DOCX tests ─────────────────────────────────────────────────


@pytest.mark.asyncio
async def test_docx_paragraphs_and_tables() -> None:
    blob = _make_docx(
        paragraphs=["Hello world", "Second paragraph"],
        table_rows=[["A1", "B1"], ["A2", "B2"]],
    )
    text, truncated = await extract_text(
        blob,
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        max_bytes=64_000,
    )
    assert "Hello world" in text
    assert "Second paragraph" in text
    assert "A1 | B1" in text
    assert "A2 | B2" in text
    assert truncated is False


@pytest.mark.asyncio
async def test_docx_byte_cap() -> None:
    blob = _make_docx(paragraphs=["x" * 200 for _ in range(20)])
    text, truncated = await extract_text(
        blob,
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        max_bytes=500,
    )
    assert truncated is True
    assert "[TRUNCATED]" in text


@pytest.mark.asyncio
async def test_docx_corrupt_bytes_raises() -> None:
    with pytest.raises(ValueError, match="AD-720a-1"):
        await extract_text(
            b"definitely not a docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            max_bytes=64_000,
        )


# ── XLSX tests ─────────────────────────────────────────────────


@pytest.mark.asyncio
async def test_xlsx_multiple_sheets() -> None:
    blob = _make_xlsx({
        "S1": [["a", "b"], ["c", "d"]],
        "S2": [["e", "f"]],
    })
    text, truncated = await extract_text(
        blob,
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        max_bytes=64_000,
    )
    assert "=== Sheet: S1 ===" in text
    assert "=== Sheet: S2 ===" in text
    assert "a | b" in text
    assert "e | f" in text
    assert truncated is False


@pytest.mark.asyncio
async def test_xlsx_row_cap() -> None:
    rows = [[f"row{i}"] for i in range(_MAX_XLSX_ROWS + 100)]
    blob = _make_xlsx({"Big": rows})
    text, truncated = await extract_text(
        blob,
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        max_bytes=10_000_000,
    )
    assert truncated is True
    assert "[TRUNCATED]" in text


@pytest.mark.asyncio
async def test_xlsx_byte_cap() -> None:
    rows = [[f"row{i}_with_some_content"] for i in range(100)]
    blob = _make_xlsx({"S": rows})
    text, truncated = await extract_text(
        blob,
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        max_bytes=200,
    )
    assert truncated is True
    assert "[TRUNCATED]" in text


# ── Dispatcher + config regression ─────────────────────────────


@pytest.mark.asyncio
async def test_dispatch_unknown_mime_still_raises_value_error() -> None:
    with pytest.raises(ValueError, match="unsupported MIME"):
        await extract_text(b"x", "application/x-unknown-type", max_bytes=100)


def test_pdf_extraction_default_off_flag() -> None:
    """Regression: AttachmentsConfig.pdf_extraction_enabled defaults to False."""
    cfg = SystemConfig()
    assert cfg.attachments.pdf_extraction_enabled is False
