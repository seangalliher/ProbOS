"""AD-755 tests for local template registry."""

from __future__ import annotations

import json
from datetime import datetime, timezone
from pathlib import Path

import pytest
from docx import Document

from probos.integrations.template_registry import (
    Template,
    TemplateRegistry,
    compute_template_ref,
    serialize_template_entry,
)


def _build_docx_bytes(text: str) -> bytes:
    doc = Document()
    doc.add_paragraph(text)
    path = Path("temp.docx")
    doc.save(str(path))
    data = path.read_bytes()
    path.unlink(missing_ok=True)
    return data


@pytest.mark.asyncio
async def test_template_load_and_list(tmp_path: Path) -> None:
    registry = TemplateRegistry(str(tmp_path))

    content = _build_docx_bytes("Hello {{name}}")
    template = Template(
        id="tmpl-1",
        name="Meeting Notes",
        file_type="docx",
        content=content,
        created_at=datetime.now(timezone.utc),
    )
    ref = compute_template_ref(content)
    (tmp_path / "blobs" / f"{ref}.bin").write_bytes(content)
    (tmp_path / "index.json").write_text(
        json.dumps([serialize_template_entry(template, ref)]),
        encoding="utf-8",
    )

    templates = await registry.list_templates()

    assert len(templates) == 1
    assert templates[0].id == "tmpl-1"
    assert templates[0].name == "Meeting Notes"


@pytest.mark.asyncio
async def test_template_instantiation_with_substitutions(tmp_path: Path) -> None:
    registry = TemplateRegistry(str(tmp_path))

    content = _build_docx_bytes("Captain: {{captain}}")
    template = Template(
        id="tmpl-2",
        name="Weekly Report",
        file_type="docx",
        content=content,
        created_at=datetime.now(timezone.utc),
    )
    ref = compute_template_ref(content)
    (tmp_path / "blobs" / f"{ref}.bin").write_bytes(content)
    (tmp_path / "index.json").write_text(
        json.dumps([serialize_template_entry(template, ref)]),
        encoding="utf-8",
    )

    output = await registry.create_from_template(
        "tmpl-2",
        substitutions={"captain": "Picard"},
        output_path=str(tmp_path / "rendered.docx"),
    )

    rendered = Document(output)
    text = "\n".join(p.text for p in rendered.paragraphs)
    assert "Picard" in text
